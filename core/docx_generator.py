import io
import os

from PIL import Image
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

SEVERITY_STYLE = {
    "CRITICAL": ("FFD1D1", "C00000"), "วิกฤต": ("FFD1D1", "C00000"),
    "WARNING": ("FFE699", "D66011"), "เตือน": ("FFE699", "D66011"),
    "NORMAL": ("D9EAD3", "38761D"), "ปกติ": ("D9EAD3", "38761D"),
    "INFORMATIONAL": ("DDEBF7", "1F4E79"), "ข้อมูลทั่วไป": ("DDEBF7", "1F4E79"),
    "EXTREME": ("F4CCCC", "990000"), "HIGH": ("FCE5CD", "A61C00"),
    "MEDIUM": ("FFF2CC", "7F6000"), "LOW": ("D9EAD3", "274E13"),
}


def _text(value):
    return "" if value is None else str(value)


def _heading(doc, text, size=11):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(_text(text))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)
    return paragraph


def _shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))


def _severity_cell(cell, value):
    style = SEVERITY_STYLE.get(_text(value).strip().upper()) or SEVERITY_STYLE.get(_text(value).strip())
    if not style:
        return
    _shade(cell, style[0])
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(style[1])


def _table(doc, headers, rows, severity_column=None):
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = _text(header)
        _shade(cell, "17365D")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row)):
            cell.text = _text(value)
            cell.margin_top = Inches(0.04)
            cell.margin_bottom = Inches(0.04)
            if row_index % 2:
                _shade(cell, "F3F6FA")
            if severity_column == index:
                _severity_cell(cell, value)
    return table


def _add_header(doc, language="en"):
    logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "logo.png")
    for section in doc.sections:
        table = section.header.add_table(rows=1, cols=2, width=Inches(7.2))
        left, right = table.rows[0].cells
        try:
            if os.path.exists(logo_path):
                left.paragraphs[0].add_run().add_picture(logo_path, height=Inches(0.3))
            else:
                raise OSError("logo unavailable")
        except Exception:
            left.text = "STUDIO OM"
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_text = "รายงานวิศวกรรม Solar O&M" if language == "th" else "Solar O&M Engineering Report"
        run = right.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.name = "Arial"


def _add_photos(doc, uploaded_files):
    for file in uploaded_files or []:
        name = _text(getattr(file, "name", "evidence"))
        if not name.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = paragraph.add_run(name)
        caption.bold = True
        try:
            file.seek(0)
            with Image.open(file) as image:
                image.thumbnail((1400, 1000), Image.Resampling.LANCZOS)
                buffer = io.BytesIO()
                image.convert("RGB").save(buffer, format="JPEG", quality=95, subsampling=0)
                buffer.seek(0)
                doc.add_paragraph().add_run().add_picture(buffer, width=Inches(6.0))
            file.seek(0)
        except Exception:
            file.seek(0)


def build_docx(report: dict, uploaded_files=None) -> io.BytesIO:
    doc = Document()
    _add_header(doc, report.get("language", "th"))
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    language = report.get("language", "th")
    labels = {
        "th": {
            "title": "รายงานวิศวกรรม Solar O&M", "executive": "สรุปผู้บริหาร", "findings": "ผลการตรวจสอบจากหลักฐาน",
            "causes": "สาเหตุราก", "damage": "ตารางความเสียหายต่อเนื่องและความเสี่ยง",
            "actions": "การแก้ไข", "spares": "อะไหล่และเครื่องมือ", "photos": "ภาคผนวก A หลักฐานภาพ",
            "summary_headers": ["โรงไฟฟ้า", "กำลังติดตั้ง", "วันที่ตรวจ", "สถานะ", "กำลังผลิตจริง", "กระแสไฟฟ้าเข้าระบบ"],
            "finding_headers": ["หมวดหมู่", "แหล่งข้อมูล", "ข้อมูลที่สังเกตได้", "การวินิจฉัยทางวิศวกรรม", "ระดับความรุนแรง"],
            "damage_headers": ["รหัส", "อุปกรณ์ที่ได้รับผลกระทบ", "รูปแบบความขัดข้อง", "ผลกระทบ", "โอกาสเกิด", "ความรุนแรง", "คะแนน", "ระดับความเสี่ยง", "ผลกระทบทางการเงิน", "ผลกระทบด้านหยุดเดินเครื่อง", "ผลกระทบด้านความปลอดภัย", "มาตรการลดความเสี่ยง", "หลักฐานอ้างอิง"],
            "cause_headers": ["ประเด็น", "รายละเอียด", "หลักฐานสนับสนุน"], "spare_headers": ["รายการ", "จำนวนแนะนำ", "วัตถุประสงค์"],
        },
        "en": {
            "title": "Solar O&M Engineering Report", "executive": "Executive Summary", "findings": "Evidence Findings",
            "causes": "Root Causes", "damage": "Consequential Damage & Risk Matrix",
            "actions": "Corrective Actions", "spares": "Spare Parts and Tools", "photos": "Appendix A Evidence Photos",
            "summary_headers": ["Plant", "Capacity", "Audit date", "Status", "Active power", "Grid current"],
            "finding_headers": ["Category", "Source", "Observed data", "Engineering diagnosis", "Severity"],
            "damage_headers": ["ID", "Affected asset", "Failure mode", "Consequence", "Likelihood", "Severity", "Score", "Risk level", "Financial impact", "Downtime impact", "Safety impact", "Mitigation", "Evidence basis"],
            "cause_headers": ["Issue", "Description", "Supporting evidence"], "spare_headers": ["Item", "Recommended quantity", "Purpose"],
        },
    }.get(language, {})
    if not labels:
        labels = {
            "title": "Solar O&M Engineering Report", "executive": "Executive Summary", "findings": "Evidence Findings",
            "causes": "Root Causes", "damage": "Consequential Damage & Risk Matrix", "actions": "Corrective Actions",
            "spares": "Spare Parts and Tools", "photos": "Appendix A Evidence Photos",
            "summary_headers": ["Plant", "Capacity", "Audit date", "Status", "Active power", "Grid current"],
            "finding_headers": ["Category", "Source", "Observed data", "Engineering diagnosis", "Severity"],
            "damage_headers": ["ID", "Affected asset", "Failure mode", "Consequence", "Likelihood", "Severity", "Score", "Risk level", "Financial impact", "Downtime impact", "Safety impact", "Mitigation", "Evidence basis"],
            "cause_headers": ["Issue", "Description", "Supporting evidence"], "spare_headers": ["Item", "Recommended quantity", "Purpose"],
        }
    summary = report.get("plant_summary", {})
    _heading(doc, f"STUDIO OM | {labels['title'].upper()}", 10)
    _heading(doc, summary.get("plant_name") or "Solar O&M Engineering Report", 16)
    _table(doc, labels["summary_headers"], [[
        summary.get("plant_name"), summary.get("rated_capacity_kw"), summary.get("audit_date"),
        summary.get("overall_status"), summary.get("active_power_kw"), summary.get("grid_current_a"),
    ]], severity_column=3)
    _heading(doc, labels["executive"])
    doc.add_paragraph(_text(report.get("executive_summary")))
    _heading(doc, labels["findings"])
    _table(doc, labels["finding_headers"], [
        [item.get("category"), item.get("source_file"), item.get("observed_data"), item.get("engineering_diagnosis"), item.get("severity")]
        for item in report.get("evidence_findings", []) if isinstance(item, dict)
    ], severity_column=4)
    _heading(doc, labels["causes"])
    _table(doc, labels["cause_headers"], [
        [item.get("issue"), item.get("description"), item.get("supporting_evidence")]
        for item in report.get("root_causes", []) if isinstance(item, dict)
    ])
    _heading(doc, labels["damage"])
    _table(doc, labels["damage_headers"], [
        [item.get("risk_id"), item.get("affected_asset"), item.get("failure_mode"), item.get("consequence"),
         item.get("likelihood"), item.get("severity"), item.get("risk_score"), item.get("risk_level"),
         item.get("financial_impact"), item.get("downtime_impact"), item.get("safety_impact"),
         item.get("mitigation"), item.get("evidence_basis")]
        for item in report.get("consequential_damage_risk_matrix", []) if isinstance(item, dict)
    ], severity_column=7)
    _heading(doc, labels["actions"])
    for action in report.get("corrective_actions", []):
        if isinstance(action, dict):
            _heading(doc, f"{action.get('step_number', '')}. {action.get('title', '')}", 10)
            for step in action.get("actions", []):
                doc.add_paragraph(_text(step), style="List Bullet")
    _heading(doc, labels["spares"])
    _table(doc, labels["spare_headers"], [
        [item.get("item_name"), item.get("recommended_qty"), item.get("purpose")]
        for item in report.get("spare_parts_tools", []) if isinstance(item, dict)
    ])
    if uploaded_files:
        _heading(doc, labels["photos"])
        _add_photos(doc, uploaded_files)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
